"""
pages/3_Resume_Rewriter.py

Self-contained Resume Rewriter.
Upload resume + JD here directly. Runs the analyzer internally,
then sends results to OpenAI for a grounded section-by-section rewrite.
Side-by-side preview with inline editing and feedback loop.
Score the rewrite when ready.
"""

from io import BytesIO
import os

import streamlit as st

from analyzer import analyze_documents
from resume_rewriter import rewrite_resume
from utils import EmptyDocumentError, FileTooLargeError, UnsupportedFileTypeError

st.set_page_config(page_title="Resume Rewriter", page_icon="✍️", layout="wide")

st.title("Resume Rewriter")
st.write(
    "Upload your resume and a job description. "
    "The analyzer scores the pair, then OpenAI rewrites your resume "
    "to close the gap — grounded strictly to your actual experience."
)

# ---------------------------------------------------------------------------
# OpenAI key check
# ---------------------------------------------------------------------------
openai_key = os.getenv("OPENAI_API_KEY", "").strip()
if not openai_key:
    try:
        openai_key = st.secrets.get("OPENAI_API_KEY", "").strip()
    except Exception:
        openai_key = ""

if not openai_key:
    st.error("OpenAI API key not found. Add it in Streamlit Cloud → Settings → Secrets.")
    st.stop()

# ---------------------------------------------------------------------------
# File uploads
# ---------------------------------------------------------------------------
col_a, col_b = st.columns(2)
with col_a:
    resume_file = st.file_uploader(
        "Upload Resume (.pdf, .docx)",
        type=["pdf", "docx"],
        key="rw_resume",
    )
with col_b:
    job_file = st.file_uploader(
        "Upload Job Description (.pdf, .docx)",
        type=["pdf", "docx"],
        key="rw_job",
    )

if resume_file:
    resume_file.seek(0)
    st.session_state["rw_resume_bytes"] = resume_file.read()
    st.session_state["rw_resume_name"] = resume_file.name
    resume_file.seek(0)

if job_file:
    job_file.seek(0)
    st.session_state["rw_job_bytes"] = job_file.read()
    st.session_state["rw_job_name"] = job_file.name
    job_file.seek(0)


class _FakeFile:
    def __init__(self, name, data):
        self.name = name
        self.size = len(data)
        self._buf = BytesIO(data)

    def read(self):
        return self._buf.read()

    def seek(self, p):
        return self._buf.seek(p)


def _make_text_file(name, text):
    return _FakeFile(name, text.encode("utf-8"))


active_resume = None
active_job = None

if st.session_state.get("rw_resume_bytes"):
    active_resume = _FakeFile(
        st.session_state["rw_resume_name"],
        st.session_state["rw_resume_bytes"],
    )
if st.session_state.get("rw_job_bytes"):
    active_job = _FakeFile(
        st.session_state["rw_job_name"],
        st.session_state["rw_job_bytes"],
    )

# ---------------------------------------------------------------------------
# Controls
# ---------------------------------------------------------------------------
ready = bool(active_resume and active_job)

user_feedback_input = st.text_area(
    "Feedback for the rewrite (optional on first run)",
    placeholder=(
        "e.g. Keep the Army bullet exactly as written. "
        "Strengthen the Computershare outcome language. "
        "Make the summary shorter and more direct."
    ),
    height=80,
    key="rw_feedback",
)

has_draft = bool(st.session_state.get("rw_draft"))
btn_label = "Generate Rewrite" if not has_draft else "Regenerate with Feedback"
run_clicked = st.button(btn_label, type="primary", disabled=not ready)

if not ready:
    st.caption("Upload both files above to enable the rewriter.")

# ---------------------------------------------------------------------------
# Run analysis + rewrite
# ---------------------------------------------------------------------------
if run_clicked and active_resume and active_job:
    with st.spinner("Analyzing resume against job description..."):
        try:
            result = analyze_documents(active_resume, active_job)
            st.session_state["rw_result"] = result
            st.session_state["rw_baseline_score"] = result["resume_strength_score"]
        except (UnsupportedFileTypeError, EmptyDocumentError, FileTooLargeError) as exc:
            st.error(str(exc))
            st.stop()
        except Exception as exc:
            st.error(f"Analysis failed: {exc}")
            st.stop()

    score = result["resume_strength_score"]
    st.info(f"Baseline score: **{score}%** — {result['fit_band']}. Generating rewrite...")

    with st.spinner("Rewriting — grounded strictly to your original experience..."):
        try:
            draft = rewrite_resume(
                resume_text=result["resume_text"],
                analysis_result=result,
                user_feedback=user_feedback_input,
            )
            st.session_state["rw_draft"] = draft
        except Exception as exc:
            st.error(f"Rewrite failed: {exc}")
            st.stop()

# ---------------------------------------------------------------------------
# Show draft
# ---------------------------------------------------------------------------
draft = st.session_state.get("rw_draft")
result = st.session_state.get("rw_result")
current_score = st.session_state.get("rw_baseline_score", 0)

if not draft or not result:
    st.stop()

resume_text = result.get("resume_text", "")
resume_lines = resume_text.splitlines()

st.success(f"Rewrite ready. Baseline: **{current_score}%**. Edit any section, then score.")

# ---------------------------------------------------------------------------
# Summary
# ---------------------------------------------------------------------------
st.divider()
st.markdown("### Professional Summary")
cl, cr = st.columns(2)

original_summary = ""
for i, line in enumerate(resume_lines):
    if "summary" in line.lower() or "profile" in line.lower():
        original_summary = "\n".join(resume_lines[i + 1 : i + 6]).strip()
        break
if not original_summary:
    original_summary = "\n".join(resume_lines[:5]).strip()

with cl:
    st.text_area("Original", value=original_summary, height=150, disabled=True, key="orig_summary")
with cr:
    edited_summary = st.text_area(
        "Rewritten (edit freely)",
        value=draft.get("summary", ""),
        height=150,
        key="edit_summary",
    )

# ---------------------------------------------------------------------------
# Experience
# ---------------------------------------------------------------------------
st.divider()
st.markdown("### Experience")

experience_drafts = draft.get("experience", [])
edited_bullets = {}

for i, role_draft in enumerate(experience_drafts):
    st.markdown(
        f"**{role_draft.get('role', '')}** — "
        f"{role_draft.get('firm', '')} "
        f"({role_draft.get('dates', '')})"
    )
    cl, cr = st.columns(2)

    firm_key = role_draft.get("firm", "").lower()[:8]
    original_block = ""
    for j, line in enumerate(resume_lines):
        if firm_key and firm_key in line.lower():
            original_block = "\n".join(resume_lines[j : j + 8]).strip()
            break

    with cl:
        st.text_area(
            "Original",
            value=original_block or "(see original resume)",
            height=180,
            disabled=True,
            key=f"orig_role_{i}",
        )
    with cr:
        bullets_text = "\n".join(f"• {b}" for b in role_draft.get("bullets", []))
        edited = st.text_area(
            "Rewritten (edit freely)",
            value=bullets_text,
            height=180,
            key=f"edit_role_{i}",
        )
        edited_bullets[i] = edited

# ---------------------------------------------------------------------------
# Skills
# ---------------------------------------------------------------------------
st.divider()
st.markdown("### Skills")
cl, cr = st.columns(2)

original_skills = ""
for j, line in enumerate(resume_lines):
    if any(w in line.lower() for w in ("skill", "expertise", "competenc")):
        original_skills = "\n".join(resume_lines[j : j + 6]).strip()
        break

with cl:
    st.text_area("Original", value=original_skills or "(see resume)", height=120, disabled=True, key="orig_skills")
with cr:
    edited_skills = st.text_area(
        "Rewritten (edit freely)",
        value=draft.get("skills", ""),
        height=120,
        key="edit_skills",
    )

# ---------------------------------------------------------------------------
# Grounding notes
# ---------------------------------------------------------------------------
st.divider()
with st.expander("Grounding notes — what was reframed and why"):
    for note in draft.get("grounding_notes", []) or ["No grounding notes returned."]:
        st.write(f"- {note}")

# ---------------------------------------------------------------------------
# Score the rewrite
# ---------------------------------------------------------------------------
st.divider()
st.markdown("### Score the Rewrite")

if st.button("Score Rewritten Resume", type="secondary"):
    parts = [edited_summary, ""]
    for i, role_draft in enumerate(experience_drafts):
        parts.append(
            f"{role_draft.get('role','')} | {role_draft.get('firm','')} | {role_draft.get('dates','')}"
        )
        for line in edited_bullets.get(i, "").splitlines():
            clean = line.strip().lstrip("•").strip()
            if clean:
                parts.append(f"• {clean}")
        parts.append("")
    parts += ["KEY SKILLS & EXPERTISE", edited_skills]

    with st.spinner("Scoring rewritten resume..."):
        try:
            jd_fake = _make_text_file(
                result.get("job_filename", "jd.txt"),
                result.get("job_text", ""),
            )
            resume_fake = _make_text_file("rewritten_resume.txt", "\n".join(parts))
            rescore = analyze_documents(resume_fake, jd_fake)
            new_score = rescore["resume_strength_score"]
            delta = new_score - current_score

            c1, c2, c3 = st.columns(3)
            with c1:
                st.metric("Original Score", f"{current_score}%")
            with c2:
                st.metric("Rewritten Score", f"{new_score}%", delta=f"{delta:+}%")
            with c3:
                st.metric("Fit Band", rescore["fit_band"])

            if new_score >= 90:
                st.success("Target reached. Resume is strongly aligned with this role.")
            elif new_score >= 80:
                st.info("Strong. Add feedback above and regenerate to push further.")
            else:
                st.warning("Still below 90%. Add feedback and regenerate.")

            if rescore.get("missing_keywords"):
                st.write("**Remaining missing keywords:**")
                for kw in rescore["missing_keywords"]:
                    st.write(f"- {kw}")

        except Exception as exc:
            st.error(f"Scoring failed: {exc}")

# ---------------------------------------------------------------------------
# Download as Word document — MD-level financial services format spec
# Calibri 9.5pt, all black, 0.5" margins, bold+underlined section headers,
# role | firm dates right-tab, page break before Key Skills
# ---------------------------------------------------------------------------
st.divider()

def build_docx(summary, experience_drafts, edited_bullets, skills, resume_text=""):
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    BLACK = RGBColor(0x00, 0x00, 0x00)
    FONT = "Calibri"
    BODY = Pt(9.5)

    lines = (resume_text or "").splitlines()
    name = lines[0].strip() if lines else "Candidate Name"
    tagline = lines[1].strip() if len(lines) > 1 else ""
    contact = lines[2].strip() if len(lines) > 2 else ""

    def _set_run(run, bold=False, italic=False, underline=False):
        run.font.name = FONT
        run.font.color.rgb = BLACK
        run.font.size = BODY
        run.bold = bold
        run.italic = italic
        run.underline = underline

    def _set_spacing(para, before=0, after=0, line=None):
        pPr = para._p.get_or_add_pPr()
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:before"), str(before))
        spacing.set(qn("w:after"), str(after))
        if line:
            spacing.set(qn("w:line"), str(line))
            spacing.set(qn("w:lineRule"), "auto")
        pPr.append(spacing)

    def _section_header(text):
        p = doc.add_paragraph()
        _set_spacing(p, before=100, after=30)
        run = p.add_run(text.upper())
        _set_run(run, bold=True, underline=True)

    def _role_line(role, firm, dates):
        p = doc.add_paragraph()
        _set_spacing(p, before=60, after=0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), WD_ALIGN_PARAGRAPH.RIGHT)
        r1 = p.add_run(role)
        _set_run(r1, bold=True)
        r2 = p.add_run(f"  |  {firm}")
        _set_run(r2, italic=True)
        r3 = p.add_run(f"\t{dates}")
        _set_run(r3)

    def _bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        _set_spacing(p, before=0, after=0, line=218)
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.14)
        _set_run(p.add_run(text))

    def _body(text, before=0, after=0):
        p = doc.add_paragraph()
        _set_spacing(p, before=before, after=after, line=218)
        _set_run(p.add_run(text))

    def _page_break():
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=0)
        br = OxmlElement("w:br")
        br.set(qn("w:type"), "page")
        p.add_run()._r.append(br)

    doc = DocxDocument()

    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(0.5)
        sec.right_margin = Inches(0.5)

    doc.styles["Normal"].paragraph_format.space_before = Pt(0)
    doc.styles["Normal"].paragraph_format.space_after = Pt(0)
    doc.styles["Normal"].font.name = FONT
    doc.styles["Normal"].font.size = BODY
    doc.styles["Normal"].font.color.rgb = BLACK

    p = doc.add_paragraph()
    _set_spacing(p, before=0, after=10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_run(p.add_run(name), bold=True)

    if tagline:
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=6)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(p.add_run(tagline), italic=True)

    if contact:
        p = doc.add_paragraph()
        _set_spacing(p, before=0, after=80)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(p.add_run(contact))

    _section_header("Professional Summary")
    _body(summary, before=30)

    _section_header("Professional Experience")
    for i, role_draft in enumerate(experience_drafts):
        _role_line(
            role_draft.get("role", ""),
            role_draft.get("firm", ""),
            role_draft.get("dates", ""),
        )

        bullets_raw = edited_bullets.get(i, "")
        for line in bullets_raw.splitlines():
            clean = line.strip().lstrip("•").strip()
            if clean:
                _bullet(clean)
        if i < len(experience_drafts) - 1:
            _body("", before=50)

    _page_break()

    _section_header("Key Skills & Expertise")
    _body(skills, before=30)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


try:
    docx_bytes = build_docx(
        summary=edited_summary,
        experience_drafts=experience_drafts,
        edited_bullets=edited_bullets,
        skills=edited_skills,
        resume_text=result.get("resume_text", ""),
    )
    st.download_button(
        label="Download Rewritten Resume (.docx)",
        data=docx_bytes,
        file_name="rewritten_resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
except Exception as exc:
    st.error(f"Word export failed: {exc}")
